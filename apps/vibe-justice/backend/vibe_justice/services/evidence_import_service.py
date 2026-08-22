"""Bounded local evidence import, validation, persistence, and extraction."""
from __future__ import annotations
import hashlib, os, shutil, warnings, zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional
from uuid import uuid4
from docx import Document
from fastapi import HTTPException, UploadFile
from PIL import Image
from pypdf import PdfReader
from sqlmodel import Session, select
from vibe_justice.api.cases import _load_case
from vibe_justice.models.evidence import EvidenceRecord, ExtractionAttempt
from vibe_justice.utils.paths import get_data_directory
from vibe_justice.utils.database import create_runtime_engine, upgrade_database
try:
    import filetype
except ImportError:  # The pinned dependency is present in production; parsers remain authoritative.
    filetype = None

CHUNK_SIZE = 1024 * 1024
MAX_FILE_BYTES = int(os.getenv("VIBE_JUSTICE_EVIDENCE_MAX_BYTES", str(25 * 1024 * 1024)))
MAX_DOCX_ENTRIES, MAX_DOCX_UNCOMPRESSED, MAX_DOCX_RATIO = 2000, 100 * 1024 * 1024, 200
MAX_IMAGE_PIXELS, MAX_PDF_PAGES = 50_000_000, 2000
ALLOWED = {".pdf", ".docx", ".txt", ".png", ".jpg", ".jpeg", ".tif", ".tiff"}
MIMES = {".pdf":"application/pdf", ".docx":"application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".txt":"text/plain", ".png":"image/png", ".jpg":"image/jpeg", ".jpeg":"image/jpeg", ".tif":"image/tiff", ".tiff":"image/tiff"}
RESERVED = {"CON", "PRN", "AUX", "NUL", *(f"COM{i}" for i in range(1,10)), *(f"LPT{i}" for i in range(1,10))}

class EvidenceImportService:
    def __init__(self):
        self.data_root = get_data_directory().resolve()
        self.staging_root = self.data_root / ".evidence-staging"; self.staging_root.mkdir(parents=True, exist_ok=True)
        upgrade_database()
        self.engine = create_runtime_engine()
        self.reconcile()

    @staticmethod
    def validate_filename(raw):
        name=(raw or "").strip()
        if not name or len(name)>255 or Path(name).name != name or any(ord(c)<32 for c in name) or name.endswith(("."," ")): raise HTTPException(400,"Invalid evidence filename")
        suffix=Path(name).suffix.lower(); stem=Path(name).stem
        if Path(stem).suffix: raise HTTPException(400,"Ambiguous double extension")
        if suffix not in ALLOWED: raise HTTPException(415,"Unsupported evidence format")
        if stem.upper() in RESERVED: raise HTTPException(400,"Reserved filename")
        return name,suffix

    def relative(self,path):
        try: return path.resolve().relative_to(self.data_root).as_posix()
        except ValueError as exc: raise RuntimeError("Evidence path escaped the data root") from exc
    def resolve(self,relative):
        path=(self.data_root/relative).resolve()
        try: path.relative_to(self.data_root)
        except ValueError as exc: raise RuntimeError("Stored evidence path is invalid") from exc
        return path

    async def stream(self,upload,staging):
        digest=hashlib.sha256(); size=0
        with staging.open("xb") as output:
            while chunk := await upload.read(CHUNK_SIZE):
                size += len(chunk)
                if size>MAX_FILE_BYTES: raise HTTPException(413,"Evidence file exceeds the size limit")
                output.write(chunk); digest.update(chunk)
            output.flush(); os.fsync(output.fileno())
        if not size: raise HTTPException(400,"Evidence file is empty")
        return size,digest.hexdigest()

    def validate_docx(self,path):
        try:
            with zipfile.ZipFile(path) as archive:
                entries=archive.infolist(); names={e.filename for e in entries}
                if len(entries)>MAX_DOCX_ENTRIES: raise HTTPException(422,"DOCX has too many entries")
                if "[Content_Types].xml" not in names or "word/document.xml" not in names: raise HTTPException(422,"Invalid DOCX package")
                total=0
                for entry in entries:
                    pure=Path(entry.filename.replace("\\","/")); total += entry.file_size
                    if pure.is_absolute() or ".." in pure.parts: raise HTTPException(422,"Unsafe DOCX package")
                    if total>MAX_DOCX_UNCOMPRESSED: raise HTTPException(422,"DOCX expanded size exceeds the limit")
                    if entry.compress_size and entry.file_size/entry.compress_size>MAX_DOCX_RATIO: raise HTTPException(422,"Suspicious DOCX compression ratio")
            Document(path)
        except HTTPException: raise
        except Exception as exc: raise HTTPException(422,"Malformed DOCX file") from exc

    def validate(self,path,suffix):
        header=path.open("rb").read(261); encrypted=False
        kind=filetype.guess(header) if filetype else None
        expected_family={".pdf":{"application/pdf"},".png":{"image/png"},".jpg":{"image/jpeg"},".jpeg":{"image/jpeg"},".tif":{"image/tiff"},".tiff":{"image/tiff"}}.get(suffix)
        if kind and expected_family and kind.mime not in expected_family:
            raise HTTPException(422,"File signature does not match extension")
        if suffix==".pdf":
            if not header.startswith(b"%PDF-"): raise HTTPException(422,"File signature does not match PDF")
            try:
                reader=PdfReader(path); encrypted=reader.is_encrypted
                if not encrypted and len(reader.pages)>MAX_PDF_PAGES: raise HTTPException(422,"PDF page count exceeds the limit")
            except HTTPException: raise
            except Exception as exc: raise HTTPException(422,"Malformed PDF file") from exc
        elif suffix==".docx":
            if not header.startswith(b"PK"): raise HTTPException(422,"File signature does not match DOCX")
            self.validate_docx(path)
        elif suffix in {".png",".jpg",".jpeg",".tif",".tiff"}:
            expected={".png":"PNG",".jpg":"JPEG",".jpeg":"JPEG",".tif":"TIFF",".tiff":"TIFF"}[suffix]
            try:
                with warnings.catch_warnings():
                    warnings.simplefilter("error"); Image.MAX_IMAGE_PIXELS=MAX_IMAGE_PIXELS
                    with Image.open(path) as image:
                        if image.format != expected: raise HTTPException(422,"Image signature does not match extension")
                        image.verify()
            except HTTPException: raise
            except Exception as exc: raise HTTPException(422,"Malformed image file") from exc
        else:
            sample=path.open("rb").read(1024*1024)
            if b"\0" in sample: raise HTTPException(422,"Text file contains binary data")
            try: sample.decode("utf-8-sig")
            except UnicodeDecodeError as exc: raise HTTPException(422,"Text file is not valid UTF-8") from exc
        return MIMES[suffix],encrypted

    async def import_upload(self,case_id,upload,source_label,received_from=None,notes=None,evidence_date=None):
        case,_=_load_case(case_id)
        if case.is_archived: raise HTTPException(409,"Archived case cannot accept evidence")
        display,suffix=self.validate_filename(upload.filename); evidence_id=str(uuid4()); staging=self.staging_root/f"{evidence_id}.upload"
        final=self.data_root/"cases"/case.case_id/"evidence"/evidence_id/"original"/f"{evidence_id}{suffix}"
        try:
            size,digest=await self.stream(upload,staging); mime,encrypted=self.validate(staging,suffix); final.parent.mkdir(parents=True,exist_ok=False)
            with Session(self.engine) as session:
                duplicate=session.exec(select(EvidenceRecord).where(EvidenceRecord.sha256==digest)).first()
                record=EvidenceRecord(evidence_id=evidence_id,case_id=case.case_id,display_filename=display,original_path=self.relative(final),byte_length=size,sha256=digest,declared_mime=upload.content_type,detected_mime=mime,detected_type=suffix.lstrip("."),source_label=source_label,received_from=received_from,notes=notes,evidence_date=evidence_date,same_content_as=duplicate.evidence_id if duplicate else None)
                session.add(record); session.commit(); os.replace(staging,final); record.status="stored"; record.updated_at=datetime.now(timezone.utc); session.add(record); session.commit()
            self.extract(case.case_id,evidence_id,encrypted); return self.get(case.case_id,evidence_id)
        except Exception:
            staging.unlink(missing_ok=True)
            # Before publish, compensating deletion is exact: remove only this
            # request's staged row. After publish, retain the staged row so a
            # restart can finish reconciliation without orphaning the original.
            if not final.is_file():
                with Session(self.engine) as session:
                    staged_record=session.get(EvidenceRecord,evidence_id)
                    if staged_record is not None:
                        session.delete(staged_record); session.commit()
            if final.parents[1].exists() and not any(path.is_file() for path in final.parents[1].rglob("*")): shutil.rmtree(final.parents[1],ignore_errors=True)
            self.reconcile()
            raise

    def get(self,case_id,evidence_id):
        with Session(self.engine) as session:
            record=session.exec(select(EvidenceRecord).where(EvidenceRecord.case_id==case_id,EvidenceRecord.evidence_id==evidence_id)).first()
            if not record: raise HTTPException(404,"Evidence not found")
            path=self.resolve(record.original_path)
            if not path.is_file(): record.status,record.error_code,record.error_message="missing","original_missing","Stored original is missing"
            elif self.hash_file(path)!=record.sha256: record.status,record.error_code,record.error_message="corrupt","integrity_mismatch","Stored original failed integrity verification"
            session.add(record); session.commit(); session.refresh(record); return record
    @staticmethod
    def hash_file(path):
        digest=hashlib.sha256()
        with path.open("rb") as source:
            while chunk:=source.read(CHUNK_SIZE): digest.update(chunk)
        return digest.hexdigest()
    def list(self,case_id):
        _load_case(case_id)
        with Session(self.engine) as session: ids=session.exec(select(EvidenceRecord.evidence_id).where(EvidenceRecord.case_id==case_id).order_by(EvidenceRecord.imported_at.desc())).all()
        return [self.get(case_id,item) for item in ids]
    def attempts(self,evidence_id):
        with Session(self.engine) as session: return list(session.exec(select(ExtractionAttempt).where(ExtractionAttempt.evidence_id==evidence_id).order_by(ExtractionAttempt.started_at.desc())).all())
    def extract(self,case_id,evidence_id,encrypted=None):
        record=self.get(case_id,evidence_id); original=self.resolve(record.original_path); before=self.hash_file(original)
        attempt=ExtractionAttempt(attempt_id=str(uuid4()),evidence_id=evidence_id,extractor_name="vibe-justice-local",extractor_version="1",status="running")
        try:
            text=""; pages=None; suffix="."+record.detected_type
            if suffix==".pdf":
                reader=PdfReader(original)
                if encrypted if encrypted is not None else reader.is_encrypted: attempt.status,attempt.error_code,attempt.error_message="encrypted","pdf_encrypted","PDF is encrypted"
                else: pages=len(reader.pages); text="\n\n".join((p.extract_text() or "") for p in reader.pages)
            elif suffix==".docx": text="\n".join(p.text for p in Document(original).paragraphs if p.text)
            elif suffix==".txt": text=original.read_text(encoding="utf-8-sig")
            else: attempt.status,attempt.error_code,attempt.error_message="unsupported","no_text_extractor","No local text extractor for this format"
            if attempt.status=="running":
                target=original.parents[1]/"extracted"/f"{attempt.attempt_id}.txt"; target.parent.mkdir(parents=True,exist_ok=True)
                with target.open("x",encoding="utf-8",newline="") as output: output.write(text); output.flush(); os.fsync(output.fileno())
                attempt.text_path=self.relative(target); attempt.text_sha256=hashlib.sha256(text.encode("utf-8")).hexdigest(); attempt.status="succeeded"
            attempt.page_count=pages
        except Exception: attempt.status,attempt.error_code,attempt.error_message="failed","extraction_failed","Local text extraction failed"
        attempt.completed_at=datetime.now(timezone.utc)
        if self.hash_file(original)!=before: raise RuntimeError("Original changed during extraction")
        with Session(self.engine) as session:
            session.add(attempt); stored=session.get(EvidenceRecord,evidence_id); stored.status={"succeeded":"ready","encrypted":"encrypted","unsupported":"unsupported","failed":"extraction_failed"}[attempt.status]; stored.updated_at=datetime.now(timezone.utc); session.add(stored); session.commit(); session.refresh(attempt)
        return attempt
    def reconcile(self):
        with Session(self.engine) as session:
            for record in session.exec(select(EvidenceRecord)).all():
                exists=self.resolve(record.original_path).is_file()
                if not exists: record.status,record.error_code,record.error_message="missing","original_missing","Stored original is missing"; session.add(record)
                elif record.status=="staged": record.status="stored"; record.updated_at=datetime.now(timezone.utc); session.add(record)
            session.commit()
